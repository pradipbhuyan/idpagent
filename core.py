import json
import re
import time
from io import BytesIO
from pathlib import Path
import tempfile

import pandas as pd
from docx import Document as DocxDocument
from langchain_openai import ChatOpenAI
from streamlit import session_state as st_state

MODEL_PRICING = {
    "gpt-4o-mini": {"input_per_1k": 0.00015, "output_per_1k": 0.0006},
    "gpt-4o": {"input_per_1k": 0.005, "output_per_1k": 0.015},
    "gpt-5": {"input_per_1k": 0.0, "output_per_1k": 0.0},
}

def ensure_metrics_state():
    if "metrics" not in st_state:
        st_state["metrics"] = {
            "tokens": 0,
            "input_tokens": 0,
            "output_tokens": 0,
            "cost": 0.0,
            "response_times": [],
            "calls": 0
        }

    if "doc_costs" not in st_state:
        st_state["doc_costs"] = {}

def get_current_metrics_snapshot():
    ensure_metrics_state()
    m = st_state["metrics"]
    return {
        "tokens": m.get("tokens", 0),
        "input_tokens": m.get("input_tokens", 0),
        "output_tokens": m.get("output_tokens", 0),
        "cost": m.get("cost", 0.0),
        "calls": m.get("calls", 0),
    }

def diff_metrics_snapshot(before, after):
    return {
        "tokens": after.get("tokens", 0) - before.get("tokens", 0),
        "input_tokens": after.get("input_tokens", 0) - before.get("input_tokens", 0),
        "output_tokens": after.get("output_tokens", 0) - before.get("output_tokens", 0),
        "cost": after.get("cost", 0.0) - before.get("cost", 0.0),
        "calls": after.get("calls", 0) - before.get("calls", 0),
    }

def get_model_pricing(model_name: str):
    return MODEL_PRICING.get(model_name, MODEL_PRICING.get("gpt-4o-mini"))

def invoke_llm_tracked(prompt: str):
    if "api_key" not in st_state:
        raise ValueError("Missing API key")

    model_name = st_state.get("model_choice", "gpt-4o-mini")
    llm = ChatOpenAI(model=model_name, temperature=0, api_key=st_state["api_key"])

    start = time.time()
    response = llm.invoke(prompt)
    duration = time.time() - start

    usage = getattr(response, "response_metadata", {}).get("token_usage", {}) or {}
    input_tokens = usage.get("prompt_tokens", 0)
    output_tokens = usage.get("completion_tokens", 0)

    if not input_tokens and not output_tokens:
        input_tokens = len(str(prompt)) // 4
        output_tokens = len(str(getattr(response, "content", ""))) // 4

    total_tokens = input_tokens + output_tokens
    pricing = get_model_pricing(model_name)
    input_cost = input_tokens * pricing["input_per_1k"] / 1000
    output_cost = output_tokens * pricing["output_per_1k"] / 1000
    total_cost = input_cost + output_cost

    ensure_metrics_state()
    m = st_state["metrics"]
    m["tokens"] += total_tokens
    m["input_tokens"] += input_tokens
    m["output_tokens"] += output_tokens
    m["cost"] += total_cost
    m["calls"] += 1
    m["response_times"].append(duration)

    doc = st_state.get("current_file") or "unknown"
    if doc not in st_state["doc_costs"]:
        st_state["doc_costs"][doc] = {"cost": 0.0, "tokens": 0}

    st_state["doc_costs"][doc]["cost"] += total_cost
    st_state["doc_costs"][doc]["tokens"] += total_tokens

    return response

def safe_json_parse(text):
    if not text:
        return {}

    text = text.strip().replace("```json", "").replace("```", "").strip()

    try:
        return json.loads(text)
    except Exception:
        pass

    try:
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if match:
            return json.loads(match.group())
    except Exception:
        pass

    try:
        text_fixed = re.sub(r",\s*}", "}", text)
        text_fixed = re.sub(r",\s*]", "]", text_fixed)
        return json.loads(text_fixed)
    except Exception:
        pass

    return {"raw_output": text}

def extract_structured_json(text, doc_type):
    clean_text = re.sub(r"[^\x00-\x7F]+", " ", text or "")
    clean_text = clean_text.replace("{", "").replace("}", "").strip()

    if "api_key" not in st_state:
        return {"error": "Missing API key"}

    if doc_type == "resume":
        prompt = f"""
You are a strict JSON generator.

Return ONLY valid JSON.
Do not add markdown.
Do not wrap in triple backticks.
Do not add explanation.

STRICT SCHEMA:
{{
  "name": "",
  "email": "",
  "phone": "",
  "location": "",
  "linkedin": "",
  "skills": [],
  "summary": "",
  "education": [
    {{
      "institution": "",
      "degree": "",
      "field_of_study": "",
      "start_date": "",
      "end_date": "",
      "graduation_date": "",
      "location": "",
      "details": []
    }}
  ],
  "experience": [
    {{
      "company": "",
      "role": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "is_current": false,
      "description": []
    }}
  ],
  "certifications": [
    {{
      "name": "",
      "issuer": "",
      "date": "",
      "expiry_date": ""
    }}
  ],
  "projects": [
    {{
      "name": "",
      "role": "",
      "start_date": "",
      "end_date": "",
      "description": []
    }}
  ]
}}

STRICT RULES:
- Extract ALL experience entries.
- Extract ALL education entries.
- Extract ALL certifications if present.
- Preserve ALL dates exactly as written in the CV.
- Never omit employment date ranges.
- Never omit education dates, graduation dates, or certification dates if present.
- Keep month/year or full date exactly as seen.
- If an entry says Present or Current, store that value in "end_date" exactly as written and set "is_current" to true.
- Do not merge multiple jobs into one job.
- Do not summarize away date fields.
- Use empty strings for missing scalar values.
- Use empty arrays for missing list values.
- description and details must always be arrays of strings.

CV TEXT:
{clean_text[:12000]}
"""
    elif doc_type == "invoice":
        prompt = f"""
You are a strict JSON extractor.

Return ONLY valid JSON.
No markdown.
No explanation.

Extract all identifiable invoice fields such as:
- vendor
- supplier
- invoice_number
- invoice_no
- invoice_date
- due_date
- currency
- subtotal
- tax
- total
- billing_address
- shipping_address
- purchase_order
- line_items

RULES:
- Preserve values exactly where possible
- Use arrays/objects where clearly present
- Do not summarize
- Do not invent unsupported fields

DOCUMENT TEXT:
{clean_text[:12000]}
"""
    elif doc_type == "ticket":
        prompt = f"""
You are a strict JSON extractor.

Return ONLY valid JSON.
No markdown.
No explanation.

Extract all identifiable travel ticket or expense fields such as:
- traveler_name
- ticket_number
- booking_reference
- airline
- from
- to
- departure_date
- return_date
- amount
- currency
- class
- trip_type

RULES:
- Preserve values exactly where possible
- Do not summarize
- Do not invent unsupported values

DOCUMENT TEXT:
{clean_text[:12000]}
"""
    else:
        return {}

    try:
        response = invoke_llm_tracked(prompt).content.strip()
        response = response.replace("```json", "").replace("```", "").strip()
        parsed = safe_json_parse(response)

        if isinstance(parsed, list):
            merged = {}
            for item in parsed:
                if isinstance(item, dict):
                    merged.update(item)
            parsed = merged if merged else {"data": parsed}

        if not isinstance(parsed, dict):
            parsed = {"data": parsed}

        if doc_type == "resume":
            if not parsed.get("name"):
                try:
                    name_prompt = f"""
Extract only the candidate's full name from this resume text.
Return only the name.
No explanation.

{clean_text[:3000]}
"""
                    fallback_name = invoke_llm_tracked(name_prompt).content.strip()
                    parsed["name"] = fallback_name
                except Exception:
                    parsed["name"] = "Candidate"

            for field in ["name", "email", "phone", "location", "linkedin", "summary"]:
                if field not in parsed or parsed[field] is None:
                    parsed[field] = ""

            for field in ["skills", "education", "experience", "certifications", "projects"]:
                if field not in parsed or parsed[field] is None:
                    parsed[field] = []

            normalized_education = []
            for edu in parsed.get("education", []):
                if isinstance(edu, dict):
                    normalized_education.append({
                        "institution": str(edu.get("institution", "") or ""),
                        "degree": str(edu.get("degree", "") or ""),
                        "field_of_study": str(edu.get("field_of_study", "") or ""),
                        "start_date": str(edu.get("start_date", "") or ""),
                        "end_date": str(edu.get("end_date", "") or ""),
                        "graduation_date": str(edu.get("graduation_date", "") or ""),
                        "location": str(edu.get("location", "") or ""),
                        "details": edu.get("details", []) if isinstance(edu.get("details", []), list) else []
                    })
            parsed["education"] = normalized_education

            normalized_experience = []
            for exp in parsed.get("experience", []):
                if isinstance(exp, dict):
                    normalized_experience.append({
                        "company": str(exp.get("company", "") or ""),
                        "role": str(exp.get("role", "") or ""),
                        "location": str(exp.get("location", "") or ""),
                        "start_date": str(exp.get("start_date", "") or ""),
                        "end_date": str(exp.get("end_date", "") or ""),
                        "is_current": bool(exp.get("is_current", False)),
                        "description": exp.get("description", []) if isinstance(exp.get("description", []), list) else []
                    })
            parsed["experience"] = normalized_experience

            normalized_certifications = []
            for cert in parsed.get("certifications", []):
                if isinstance(cert, dict):
                    normalized_certifications.append({
                        "name": str(cert.get("name", "") or ""),
                        "issuer": str(cert.get("issuer", "") or ""),
                        "date": str(cert.get("date", "") or ""),
                        "expiry_date": str(cert.get("expiry_date", "") or "")
                    })
            parsed["certifications"] = normalized_certifications

            normalized_projects = []
            for proj in parsed.get("projects", []):
                if isinstance(proj, dict):
                    normalized_projects.append({
                        "name": str(proj.get("name", "") or ""),
                        "role": str(proj.get("role", "") or ""),
                        "start_date": str(proj.get("start_date", "") or ""),
                        "end_date": str(proj.get("end_date", "") or ""),
                        "description": proj.get("description", []) if isinstance(proj.get("description", []), list) else []
                    })
            parsed["projects"] = normalized_projects

            if parsed.get("name"):
                parsed["name"] = parsed["name"].strip().title()

        return parsed

    except Exception as e:
        return {"error": "LLM request failed", "details": str(e)[:300]}

def generate_resume_summary(data):
    if "api_key" not in st_state:
        return "Summary not available"

    prompt = f"""
Create a professional resume summary in plain text.

STRICT RULES:
- No markdown
- No * or **
- Plain text only
- 4 to 8 concise bullet-style lines using simple hyphen prefixes
- Mention total profile strengths, major domains, and seniority
- Preserve important date context where relevant
- Do not invent dates
- Do not remove date references if they are important to career continuity
- Do not rewrite exact extracted timelines in a conflicting way

CANDIDATE DATA:
{json.dumps(data, ensure_ascii=False)}
"""
    try:
        return invoke_llm_tracked(prompt).content.strip()
    except Exception:
        return "Summary not available"

def build_resume(data, template_file):
    def safe_str(value):
        return "" if value is None else str(value)

    def format_date_range(start_date, end_date):
        start_date = safe_str(start_date).strip()
        end_date = safe_str(end_date).strip()
        if start_date and end_date:
            return f"{start_date} - {end_date}"
        if start_date:
            return start_date
        if end_date:
            return end_date
        return ""

    def format_skills(skills):
        if not isinstance(skills, list) or not skills:
            return ""
        return ", ".join(str(s).strip() for s in skills if str(s).strip())

    def format_experience(experience):
        if not isinstance(experience, list) or not experience:
            return ""
        lines = []
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            role = safe_str(exp.get("role")).strip()
            company = safe_str(exp.get("company")).strip()
            location = safe_str(exp.get("location")).strip()
            start_date = safe_str(exp.get("start_date")).strip()
            end_date = safe_str(exp.get("end_date")).strip()
            desc = exp.get("description", [])

            header_parts = []
            title_part = " - ".join([p for p in [role, company] if p])
            if title_part:
                header_parts.append(title_part)
            date_range = format_date_range(start_date, end_date)
            if date_range:
                header_parts.append(f"({date_range})")
            if location:
                header_parts.append(location)
            if header_parts:
                lines.append(" ".join(header_parts))

            if isinstance(desc, list):
                for item in desc:
                    item = safe_str(item).strip()
                    if item:
                        lines.append(f"- {item}")
            lines.append("")
        return "\n".join(lines).strip()

    def format_education(education):
        if not isinstance(education, list) or not education:
            return ""
        lines = []
        for edu in education:
            if not isinstance(edu, dict):
                continue
            institution = safe_str(edu.get("institution")).strip()
            degree = safe_str(edu.get("degree")).strip()
            field = safe_str(edu.get("field_of_study")).strip()
            location = safe_str(edu.get("location")).strip()
            start_date = safe_str(edu.get("start_date")).strip()
            end_date = safe_str(edu.get("end_date")).strip()
            graduation_date = safe_str(edu.get("graduation_date")).strip()
            details = edu.get("details", [])

            first_line_parts = []
            degree_part = ", ".join([p for p in [degree, field] if p])
            if degree_part:
                first_line_parts.append(degree_part)
            if institution:
                first_line_parts.append(institution)
            if first_line_parts:
                lines.append(" - ".join(first_line_parts))

            date_text = graduation_date if graduation_date else format_date_range(start_date, end_date)
            second_line_parts = []
            if date_text:
                second_line_parts.append(date_text)
            if location:
                second_line_parts.append(location)
            if second_line_parts:
                lines.append(", ".join(second_line_parts))

            if isinstance(details, list):
                for item in details:
                    item = safe_str(item).strip()
                    if item:
                        lines.append(f"- {item}")
            lines.append("")
        return "\n".join(lines).strip()

    def format_certifications(certifications):
        if not isinstance(certifications, list) or not certifications:
            return ""
        lines = []
        for cert in certifications:
            if not isinstance(cert, dict):
                continue
            name = safe_str(cert.get("name")).strip()
            issuer = safe_str(cert.get("issuer")).strip()
            date = safe_str(cert.get("date")).strip()
            expiry_date = safe_str(cert.get("expiry_date")).strip()

            line_parts = []
            if name:
                line_parts.append(name)
            if issuer:
                line_parts.append(issuer)
            line = " - ".join(line_parts)

            date_bits = []
            if date:
                date_bits.append(f"Issued: {date}")
            if expiry_date:
                date_bits.append(f"Expires: {expiry_date}")

            if date_bits:
                line = f"{line} ({', '.join(date_bits)})" if line else ", ".join(date_bits)
            if line:
                lines.append(line)
        return "\n".join(lines).strip()

    def format_projects(projects):
        if not isinstance(projects, list) or not projects:
            return ""
        lines = []
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            name = safe_str(proj.get("name")).strip()
            role = safe_str(proj.get("role")).strip()
            start_date = safe_str(proj.get("start_date")).strip()
            end_date = safe_str(proj.get("end_date")).strip()
            description = proj.get("description", [])

            header_parts = []
            title = " - ".join([p for p in [name, role] if p])
            if title:
                header_parts.append(title)
            date_range = format_date_range(start_date, end_date)
            if date_range:
                header_parts.append(f"({date_range})")
            if header_parts:
                lines.append(" ".join(header_parts))

            if isinstance(description, list):
                for item in description:
                    item = safe_str(item).strip()
                    if item:
                        lines.append(f"- {item}")
            lines.append("")
        return "\n".join(lines).strip()

    def replace_placeholders_in_paragraph(paragraph, placeholders):
        for key, value in placeholders.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    def replace_placeholders(doc, placeholders):
        for para in doc.paragraphs:
            replace_placeholders_in_paragraph(para, placeholders)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_placeholders_in_paragraph(para, placeholders)

        for section in doc.sections:
            for para in section.header.paragraphs:
                replace_placeholders_in_paragraph(para, placeholders)
            for para in section.footer.paragraphs:
                replace_placeholders_in_paragraph(para, placeholders)

    summary = generate_resume_summary(data)

    if not template_file:
        raise ValueError("No template file provided")

    try:
        if isinstance(template_file, bytes):
            doc = DocxDocument(BytesIO(template_file))
        elif hasattr(template_file, "read"):
            content = template_file.read()
            if not content:
                raise ValueError("Template file is empty")
            if hasattr(template_file, "seek"):
                template_file.seek(0)
            doc = DocxDocument(BytesIO(content))
        elif isinstance(template_file, str):
            doc = DocxDocument(template_file)
        else:
            raise TypeError(f"Unsupported template_file type: {type(template_file)}")
    except Exception as e:
        raise RuntimeError(f"Template load failed: {e}")

    placeholders = {
        "{{name}}": safe_str(data.get("name", "")),
        "{{email}}": safe_str(data.get("email", "")),
        "{{phone}}": safe_str(data.get("phone", "")),
        "{{location}}": safe_str(data.get("location", "")),
        "{{linkedin}}": safe_str(data.get("linkedin", "")),
        "{{summary}}": safe_str(summary),
        "{{skills}}": format_skills(data.get("skills", [])),
        "{{experience}}": format_experience(data.get("experience", [])),
        "{{education}}": format_education(data.get("education", [])),
        "{{certifications}}": format_certifications(data.get("certifications", [])),
        "{{projects}}": format_projects(data.get("projects", [])),
    }

    replace_placeholders(doc, placeholders)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def save_temp_file(uploaded_file):
    suffix = Path(uploaded_file.name).suffix
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getvalue())
        return tmp.name

def detect_document_type(text):
    if "api_key" not in st_state:
        return "other"

    prompt = f"""
Classify document into ONE label:

resume
invoice
receipt
report
ticket
other

STRICT RULES:
- Return ONLY one word
- No explanation
- No sentence

{text[:2000]}
"""
    try:
        raw = invoke_llm_tracked(prompt).content.lower().strip()
    except Exception:
        return "other"

    labels = ["resume", "invoice", "receipt", "report", "ticket", "other"]
    for label in labels:
        if label == raw:
            return label
    for label in labels:
        if label in raw:
            return label
    return "other"

def json_to_kv_dataframe(data):
    rows = []

    def flatten(prefix, obj):
        if isinstance(obj, dict):
            for k, v in obj.items():
                flatten(f"{prefix}.{k}" if prefix else k, v)
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                flatten(f"{prefix}[{i}]", item)
        else:
            rows.append({
                "Field": prefix,
                "Value": json.dumps(obj) if isinstance(obj, (dict, list)) else str(obj)
            })

    flatten("", data if data is not None else {})
    return pd.DataFrame(rows)

def generate_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="data")
    return output.getvalue()

from datetime import datetime
import uuid

def send_to_concur(doc_type, data, mode="mock"):
    """
    Mock Concur sender that looks like a real API integration response.
    Replace the `mode != "mock"` block later with actual Concur API calls.
    """

    payload = {
        "type": doc_type,
        "data": data,
    }

    if doc_type == "invoice":
        try:
            payload["line_items"] = json_to_kv_dataframe(data).to_dict(orient="records")
        except Exception:
            payload["line_items"] = []

    now_utc = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    short_id = uuid.uuid4().hex[:8].upper()
    batch_id = f"CCB-{datetime.utcnow().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"
    endpoint = "Expense Entry Import API" if doc_type == "invoice" else "Travel Request / Expense Entry API"

    if mode == "mock":
        return {
            "status": "submitted",
            "mode": "mock",
            "message": f"{doc_type.title()} submitted to Concur mock gateway",
            "submission_id": f"SUB-{short_id}",
            "batch_id": batch_id,
            "document_id": f"{doc_type[:3].upper()}-{uuid.uuid4().hex[:10].upper()}",
            "submitted_at": now_utc,
            "endpoint": endpoint,
            "processing_state": "Queued for downstream validation",
            "next_status": "Expected to transition to Accepted or Rejected after validation",
            "payload": payload
        }

    # Replace this later with real Concur API request/response handling
    return {
        "status": "submitted",
        "mode": "real",
        "message": f"{doc_type.title()} submitted to Concur",
        "submission_id": f"SUB-{short_id}",
        "batch_id": batch_id,
        "document_id": f"{doc_type[:3].upper()}-{uuid.uuid4().hex[:10].upper()}",
        "submitted_at": now_utc,
        "endpoint": endpoint,
        "processing_state": "Accepted by Concur endpoint",
        "next_status": "Awaiting downstream processing",
        "payload": payload
    }

def confidence_label(score):
    if score >= 0.85:
        return "High"
    if score >= 0.6:
        return "Medium"
    return "Low"


def build_confidence_map(data, doc_type):
    if not isinstance(data, dict):
        return {}

    def score_scalar(value, strong=False):
        if value in [None, "", [], {}]:
            return {
                "score": 0.2,
                "label": "Low",
                "reason": "Missing or empty field"
            }

        if strong:
            score = 0.9
            reason = "Looks like an explicit field match"
        else:
            score = 0.7
            reason = "Extracted successfully but may need review"

        return {
            "score": score,
            "label": confidence_label(score),
            "reason": reason
        }

    confidence = {}

    if doc_type == "invoice":
        for field in ["vendor", "invoice_number", "invoice_date", "total", "currency", "due_date"]:
            val = data.get(field) or data.get(field.replace("invoice_number", "invoice_no"))
            confidence[field] = score_scalar(val, strong=field in ["invoice_number", "total"])

    elif doc_type == "ticket":
        for field in ["traveler_name", "ticket_number", "airline", "from", "to", "departure_date", "amount"]:
            confidence[field] = score_scalar(data.get(field), strong=field in ["ticket_number", "departure_date"])

    elif doc_type == "resume":
        for field in ["name", "email", "phone", "location", "summary"]:
            confidence[field] = score_scalar(data.get(field), strong=field in ["name", "email"])

        confidence["experience"] = score_scalar(data.get("experience"), strong=True)
        confidence["education"] = score_scalar(data.get("education"), strong=True)

    return confidence


def validate_document_data(data, doc_type):
    issues = []
    warnings = []

    if not isinstance(data, dict):
        return {
            "passed": False,
            "issues": ["No structured data available"],
            "warnings": []
        }

    if doc_type == "invoice":
        if not (data.get("vendor") or data.get("supplier")):
            issues.append("Vendor is missing")
        if not (data.get("invoice_number") or data.get("invoice_no")):
            issues.append("Invoice number is missing")
        if not data.get("invoice_date"):
            issues.append("Invoice date is missing")
        if not data.get("total"):
            issues.append("Total amount is missing")

    elif doc_type == "ticket":
        if not data.get("traveler_name"):
            issues.append("Traveler name is missing")
        if not data.get("ticket_number"):
            issues.append("Ticket number is missing")
        if not data.get("from") or not data.get("to"):
            issues.append("Route is incomplete")
        if not data.get("departure_date"):
            issues.append("Departure date is missing")
        if not data.get("amount"):
            warnings.append("Amount is missing")

    elif doc_type == "resume":
        if not data.get("name"):
            issues.append("Candidate name is missing")
        if not data.get("experience"):
            issues.append("Experience section is missing")
        if not data.get("education"):
            warnings.append("Education section is missing")
        if not data.get("skills"):
            warnings.append("Skills section is missing")

    return {
        "passed": len(issues) == 0,
        "issues": issues,
        "warnings": warnings,
    }
